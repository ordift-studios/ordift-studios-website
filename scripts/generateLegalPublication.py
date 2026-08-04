#!/usr/bin/env python3
"""
Generates the four Publication Version formats (PDF, DOCX, HTML, MD) for
an Ordift Studios Enterprise Legal Series document, from the same JSON
export produced by scripts/exportLegalDocumentJson.ts — single source
of truth stays in src/lib/legal/documents/, this script never invents
or edits wording.

Usage: python3 scripts/generateLegalPublication.py <json-path> <out-dir>

The letterhead (brand/legal-publications/os-letterhead.jpg) is used
unaltered: never resized/cropped/recompressed relative to its source
file, and the PDF/DOCX page size is set to its own exact pixel aspect
ratio so it fills the page with zero stretch distortion in either
direction.
"""
import sys
import json
import os
from datetime import datetime

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LETTERHEAD_PATH = os.path.join(REPO_ROOT, "brand/legal-publications/os-letterhead.jpg")

# Measured directly from the letterhead's own pixels (see session notes):
# logo occupies rows 47-228 of 1385 (top ~16.5%), footer bar occupies
# rows 1239-1362 of 1385 (bottom ~10.5%). Page geometry below reserves
# safe zones clear of both, with a buffer, rather than guessing margins.
LETTERHEAD_PX = (1055, 1385)
PAGE_HEIGHT_PT = 792.0  # 11in, points
PAGE_WIDTH_PT = PAGE_HEIGHT_PT * (LETTERHEAD_PX[0] / LETTERHEAD_PX[1])
TOP_MARGIN_PT = 150.0
BOTTOM_MARGIN_PT = 105.0
SIDE_MARGIN_PT = 56.0


def load_doc(json_path):
    with open(json_path, "r") as f:
        return json.load(f)


def fmt_date(iso):
    return datetime.strptime(iso, "%Y-%m-%d").strftime("%-d %B %Y")


# ---------------------------------------------------------------------
# PDF (reportlab)
# ---------------------------------------------------------------------
def generate_pdf(doc, out_path):
    from reportlab.lib.pagesizes import portrait
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        BaseDocTemplate,
        PageTemplate,
        Frame,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        ListFlowable,
        ListItem,
        PageBreak,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    control = doc["control"]
    sections = doc["sections"]
    definitions = doc["definitions"]

    INK = HexColor("#1A1D24")
    INK_MUTED = HexColor("#5B5F6B")
    GOLD = HexColor("#8E7736")

    pagesize = (PAGE_WIDTH_PT, PAGE_HEIGHT_PT)

    def draw_letterhead_and_chrome(canvas_obj, doc_template):
        canvas_obj.saveState()
        canvas_obj.drawImage(
            LETTERHEAD_PATH, 0, 0, width=PAGE_WIDTH_PT, height=PAGE_HEIGHT_PT,
            preserveAspectRatio=False, mask="auto",
        )
        # Running header (only from page 2 onward — cover page stays clean)
        if canvas_obj.getPageNumber() > 1:
            canvas_obj.setFont("Helvetica", 7.5)
            canvas_obj.setFillColor(INK_MUTED)
            canvas_obj.drawRightString(
                PAGE_WIDTH_PT - SIDE_MARGIN_PT, PAGE_HEIGHT_PT - 138,
                f"{control['documentCode']} — {control['documentTitle']}",
            )
            canvas_obj.setFont("Helvetica", 7.5)
            canvas_obj.drawCentredString(
                PAGE_WIDTH_PT / 2, 93,
                f"Page {canvas_obj.getPageNumber()} · Uncontrolled when printed or downloaded · © {fmt_date(control['effectiveDate'])[-4:]} Ordift Studios",
            )
        canvas_obj.restoreState()

    frame = Frame(
        SIDE_MARGIN_PT, BOTTOM_MARGIN_PT,
        PAGE_WIDTH_PT - 2 * SIDE_MARGIN_PT, PAGE_HEIGHT_PT - TOP_MARGIN_PT - BOTTOM_MARGIN_PT,
        id="content",
    )

    class LegalDocTemplate(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                style_name = flowable.style.name
                text = flowable.getPlainText()
                if style_name == "SectionHeading":
                    self.notify("TOCEntry", (0, text, self.page))
                elif style_name == "SubHeading":
                    self.notify("TOCEntry", (1, text, self.page))

    pdf_doc = LegalDocTemplate(
        out_path, pagesize=pagesize,
        topMargin=TOP_MARGIN_PT, bottomMargin=BOTTOM_MARGIN_PT,
        leftMargin=SIDE_MARGIN_PT, rightMargin=SIDE_MARGIN_PT,
        title=control["documentTitle"], author="Ordift Studios",
    )
    pdf_doc.addPageTemplates([PageTemplate(id="main", frames=[frame], onPage=draw_letterhead_and_chrome)])

    base = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=base["Normal"], fontName="Helvetica", fontSize=9.5, leading=14, textColor=INK_MUTED, spaceAfter=8)
    cover_eyebrow = ParagraphStyle("CoverEyebrow", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=14, textColor=GOLD, alignment=TA_CENTER, spaceAfter=6)
    cover_title = ParagraphStyle("CoverTitle", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=30, leading=36, textColor=INK, alignment=TA_CENTER, spaceAfter=14)
    cover_meta = ParagraphStyle("CoverMeta", parent=base["Normal"], fontName="Helvetica", fontSize=10.5, leading=17, textColor=INK_MUTED, alignment=TA_CENTER)
    page_heading = ParagraphStyle("PageHeading", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=15, leading=19, textColor=INK, spaceAfter=12)
    section_heading = ParagraphStyle("SectionHeading", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=13, leading=17, textColor=INK, spaceBefore=14, spaceAfter=6)
    sub_heading = ParagraphStyle("SubHeading", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=10.5, leading=15, textColor=INK, spaceBefore=10, spaceAfter=4)
    plain_sub = ParagraphStyle("PlainSub", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=9.5, leading=14, textColor=INK, spaceBefore=6, spaceAfter=2)
    toc_style1 = ParagraphStyle("TOC1", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=10, leading=15, textColor=INK)
    toc_style2 = ParagraphStyle("TOC2", parent=base["Normal"], fontName="Helvetica", fontSize=9, leading=13, textColor=INK_MUTED, leftIndent=14)
    dc_label = ParagraphStyle("DCLabel", parent=base["Normal"], fontName="Helvetica-Bold", fontSize=8.5, textColor=INK_MUTED)
    dc_value = ParagraphStyle("DCValue", parent=base["Normal"], fontName="Helvetica", fontSize=9, textColor=INK)

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("“", "&#8220;").replace("”", "&#8221;").replace("’", "&#8217;")

    import re as _re

    def link_terms(text, defs):
        """Whole-word, case-sensitive: wraps each defined-term match in
        an internal PDF link to its `<a name="def-{id}"/>` anchor in the
        Definitions section (reportlab paragraph mini-markup — no extra
        canvas bookmark calls needed). Runs on already-escaped text, so
        term matching happens against the escaped form (safe: none of
        the defined terms contain characters `esc()` rewrites)."""
        if not defs:
            return text
        ordered = sorted(defs, key=lambda d: -len(d["term"]))
        pattern = _re.compile(r"\b(" + "|".join(_re.escape(d["term"]) for d in ordered) + r")\b")
        by_term = {d["term"]: d["id"] for d in ordered}

        def _sub(m):
            return f'<a href="#def-{by_term[m.group(0)]}" color="#8E7736">{m.group(0)}</a>'

        return pattern.sub(_sub, text)

    # `.` reads naturally after numeric section numbers ("5.") but not
    # after word-based appendix labels ("Appendix A.") — same rule as
    # the website's LegalSection.tsx, kept in sync deliberately.
    NUMERIC_SECTION = _re.compile(r"^[\dA-Z.]+$")

    def number_suffix(number):
        return "." if NUMERIC_SECTION.match(number) else " —"

    story = []

    # --- Cover page ---
    story.append(Spacer(1, 90))
    story.append(Paragraph(esc(control["publicationSeries"]), cover_eyebrow))
    story.append(Paragraph(esc(control["documentTitle"]), cover_title))
    story.append(Paragraph(
        f"{esc(control['documentCode'])} &nbsp;·&nbsp; Version {esc(control['version'])} &nbsp;·&nbsp; {control['status'].capitalize()}",
        cover_meta,
    ))
    story.append(Paragraph(f"Effective {fmt_date(control['effectiveDate'])}", cover_meta))
    story.append(Spacer(1, 30))
    story.append(Paragraph(f"Classification: {control['classification'].capitalize()}", cover_meta))
    story.append(PageBreak())

    # --- Document Control page ---
    story.append(Paragraph("Document Control", page_heading))
    dc_rows = [
        ["Document Title", control["documentTitle"]],
        ["Document Code", control["documentCode"]],
        ["Publication Series", control["publicationSeries"]],
        ["Version", control["version"]],
        ["Status", control["status"].capitalize()],
        ["Classification", control["classification"].capitalize()],
        ["Effective Date", fmt_date(control["effectiveDate"])],
        ["Last Updated", fmt_date(control["lastUpdated"])],
        ["Review Cycle", control["reviewCycle"]],
        ["Document Owner", control["documentOwner"]],
        ["Prepared By", control["preparedBy"]],
        ["Approved By", control["approvedBy"]],
        ["Related Documents", "; ".join(f"{(d['code'] + ' ') if d['code'] else ''}{d['title']}" for d in control["relatedDocuments"])],
    ]
    dc_table_data = [[Paragraph(esc(k), dc_label), Paragraph(esc(v), dc_value)] for k, v in dc_rows]
    dc_table = Table(dc_table_data, colWidths=[130, PAGE_WIDTH_PT - 2 * SIDE_MARGIN_PT - 130])
    dc_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, HexColor("#E5E3DE")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(dc_table)
    story.append(Spacer(1, 14))
    story.append(Paragraph(esc(control["controlledDocumentNotice"]), body))
    story.append(PageBreak())

    # --- Revision History (running change log — 2026-08-04 direction:
    # a real append-only log, not a single synthesized "initial issue" row) ---
    story.append(Paragraph("Revision History", page_heading))
    rev_data = [[Paragraph("Version", dc_label), Paragraph("Date", dc_label), Paragraph("Description", dc_label), Paragraph("Author", dc_label)]]
    for entry in control["changeLog"]:
        rev_data.append([
            Paragraph(esc(entry["version"]), dc_value), Paragraph(fmt_date(entry["date"]), dc_value),
            Paragraph(esc(entry["description"]), dc_value), Paragraph(esc(entry["author"]), dc_value),
        ])
    rev_table = Table(rev_data, colWidths=[45, 75, 245, 85])
    rev_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LINEBELOW", (0, 0), (-1, 0), 1, INK),
        ("LINEBELOW", (0, 1), (-1, 1), 0.5, HexColor("#E5E3DE")),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(rev_table)
    story.append(PageBreak())

    # --- Table of Contents ---
    story.append(Paragraph("Table of Contents", page_heading))
    toc = TableOfContents()
    toc.levelStyles = [toc_style1, toc_style2]
    story.append(toc)
    story.append(PageBreak())

    # --- Body sections ---
    for sec in sections:
        style = section_heading if sec["level"] == 1 else sub_heading
        story.append(Paragraph(f"{esc(sec['number'])}{number_suffix(sec['number'])} {esc(sec['heading'])}", style))
        # No self-links inside the Definitions section itself.
        linkable_defs = [] if sec["id"] == "definitions" else definitions
        for node in sec["content"]:
            if node["type"] == "paragraph":
                story.append(Paragraph(link_terms(esc(node["text"]), linkable_defs), body))
            elif node["type"] == "subheading":
                story.append(Paragraph(esc(node["text"]), plain_sub))
            elif node["type"] == "list":
                items = [ListItem(Paragraph(link_terms(esc(t), linkable_defs), body), leftIndent=6) for t in node["items"]]
                story.append(ListFlowable(items, bulletType="bullet" if not node.get("ordered") else "1", start="circle", leftIndent=14, bulletFontSize=7))
            elif node["type"] == "table":
                headers = [Paragraph(esc(h), dc_label) for h in node["headers"]]
                rows = [[Paragraph(esc(c), body) for c in row] for row in node["rows"]]
                col_w = (PAGE_WIDTH_PT - 2 * SIDE_MARGIN_PT) / len(node["headers"])
                t = Table([headers] + rows, colWidths=[col_w] * len(node["headers"]))
                t.setStyle(TableStyle([
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LINEBELOW", (0, 0), (-1, 0), 1, INK),
                    ("LINEBELOW", (0, 1), (-1, -1), 0.5, HexColor("#E5E3DE")),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]))
                story.append(Spacer(1, 4))
                story.append(t)
                story.append(Spacer(1, 4))
        if sec["id"] == "definitions" and definitions:
            def_items = []
            for d in definitions:
                def_items.append(Paragraph(f'<a name="def-{d["id"]}"/><b>“{esc(d["term"])}”</b> {esc(d["definition"])}', body))
            story.extend(def_items)

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        f"© {fmt_date(control['effectiveDate'])[-4:]} Ordift Studios. All rights reserved. "
        f"Document {control['documentCode']}, Version {control['version']}.",
        ParagraphStyle("Copyright", parent=base["Normal"], fontName="Helvetica", fontSize=8, textColor=INK_MUTED),
    ))

    pdf_doc.multiBuild(story)


# ---------------------------------------------------------------------
# DOCX (python-docx + raw OOXML for the full-page background image)
# ---------------------------------------------------------------------
def generate_docx(doc, out_path):
    from docx import Document
    from docx.shared import Pt, Emu, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    control = doc["control"]
    sections = doc["sections"]
    definitions = doc["definitions"]

    page_w_emu = Emu(int(PAGE_WIDTH_PT / 72 * 914400))
    page_h_emu = Emu(int(PAGE_HEIGHT_PT / 72 * 914400))

    d = Document()
    section = d.sections[0]
    section.page_width = page_w_emu
    section.page_height = page_h_emu
    section.top_margin = Pt(TOP_MARGIN_PT)
    section.bottom_margin = Pt(BOTTOM_MARGIN_PT)
    section.left_margin = Pt(SIDE_MARGIN_PT)
    section.right_margin = Pt(SIDE_MARGIN_PT)

    def add_background_image(sect):
        """Anchors the letterhead behind the text, pinned to the page
        origin at full page size — inserted once per section header so
        it repeats on every page without being resizable/movable by a
        casual edit, and without altering the source image itself."""
        header = sect.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        run.add_picture(LETTERHEAD_PATH, width=page_w_emu, height=page_h_emu)
        inline = run._element.findall(".//" + qn("wp:inline"))
        if not inline:
            return
        inline_el = inline[0]
        graphic = inline_el.find(qn("a:graphic"))
        extent = inline_el.find(qn("wp:extent"))
        docpr = inline_el.find(qn("wp:docPr"))

        anchor = OxmlElement("wp:anchor")
        anchor.set("distT", "0"); anchor.set("distB", "0"); anchor.set("distL", "0"); anchor.set("distR", "0")
        anchor.set("simplePos", "0"); anchor.set("relativeHeight", "1")
        anchor.set("behindDoc", "1"); anchor.set("locked", "1")
        anchor.set("layoutInCell", "1"); anchor.set("allowOverlap", "1")

        simple_pos = OxmlElement("wp:simplePos"); simple_pos.set("x", "0"); simple_pos.set("y", "0")
        pos_h = OxmlElement("wp:positionH"); pos_h.set("relativeFrom", "page")
        pos_h_off = OxmlElement("wp:posOffset"); pos_h_off.text = "0"
        pos_h.append(pos_h_off)
        pos_v = OxmlElement("wp:positionV"); pos_v.set("relativeFrom", "page")
        pos_v_off = OxmlElement("wp:posOffset"); pos_v_off.text = "0"
        pos_v.append(pos_v_off)
        wrap_none = OxmlElement("wp:wrapNone")

        anchor.append(simple_pos)
        anchor.append(pos_h)
        anchor.append(pos_v)
        anchor.append(extent)
        anchor.append(wrap_none)
        anchor.append(docpr)
        graphic_frame_locks_parent = inline_el.find(qn("wp:cNvGraphicFramePr"))
        if graphic_frame_locks_parent is not None:
            anchor.append(graphic_frame_locks_parent)
        anchor.append(graphic)

        parent = inline_el.getparent()
        parent.replace(inline_el, anchor)

    add_background_image(section)

    # Running footer: page number + doc code
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""
    run = fp.add_run(f"{control['documentCode']} — {control['documentTitle']} · Uncontrolled when printed or downloaded · Page ")
    run.font.size = Pt(7.5)
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    run2 = fp.add_run(); run2._r.append(fld_begin)
    run3 = fp.add_run(); run3._r.append(instr)
    run4 = fp.add_run(); run4._r.append(fld_end)
    for r in (run2, run3, run4):
        r.font.size = Pt(7.5)

    styles = d.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    def add_field(paragraph, field_code):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field_code
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        r2 = paragraph.add_run(); r2._r.append(instr)
        r3 = paragraph.add_run(); r3._r.append(fld_sep)
        r4 = paragraph.add_run(); r4._r.append(fld_end)

    # --- Cover page ---
    for _ in range(6):
        d.add_paragraph()
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(control["publicationSeries"]); run.bold = True; run.font.size = Pt(11)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(control["documentTitle"]); run.bold = True; run.font.size = Pt(30)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{control['documentCode']}  ·  Version {control['version']}  ·  {control['status'].capitalize()}")
    run.font.size = Pt(11)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Effective {fmt_date(control['effectiveDate'])}"); run.font.size = Pt(11)
    d.add_page_break()

    # --- Document Control ---
    h = d.add_heading("Document Control", level=1)
    dc_rows = [
        ("Document Title", control["documentTitle"]),
        ("Document Code", control["documentCode"]),
        ("Publication Series", control["publicationSeries"]),
        ("Version", control["version"]),
        ("Status", control["status"].capitalize()),
        ("Classification", control["classification"].capitalize()),
        ("Effective Date", fmt_date(control["effectiveDate"])),
        ("Last Updated", fmt_date(control["lastUpdated"])),
        ("Review Cycle", control["reviewCycle"]),
        ("Document Owner", control["documentOwner"]),
        ("Prepared By", control["preparedBy"]),
        ("Approved By", control["approvedBy"]),
        ("Related Documents", "; ".join(f"{(rd['code'] + ' ') if rd['code'] else ''}{rd['title']}" for rd in control["relatedDocuments"])),
    ]
    table = d.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for k, v in dc_rows:
        row = table.add_row().cells
        row[0].text = k
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        row[1].text = v
        row[1].paragraphs[0].runs[0].font.size = Pt(9)
    d.add_paragraph()
    d.add_paragraph(control["controlledDocumentNotice"]).runs[0].font.size = Pt(9)
    d.add_page_break()

    # --- Revision History (running change log) ---
    d.add_heading("Revision History", level=1)
    rev_table = d.add_table(rows=1, cols=4)
    rev_table.style = "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in styles] else None
    hdr = rev_table.rows[0].cells
    for i, label in enumerate(["Version", "Date", "Description", "Author"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True
    for entry in control["changeLog"]:
        row = rev_table.add_row().cells
        row[0].text = entry["version"]
        row[1].text = fmt_date(entry["date"])
        row[2].text = entry["description"]
        row[3].text = entry["author"]
    d.add_page_break()

    # --- Table of Contents (native Word TOC field — auto-populates on open) ---
    d.add_heading("Table of Contents", level=1)
    toc_p = d.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-2" \\h \\z \\u')
    note = d.add_paragraph()
    note_run = note.add_run("(Right-click and choose “Update Field” if section titles do not appear.)")
    note_run.italic = True
    note_run.font.size = Pt(8)
    d.add_page_break()

    # --- Body sections ---
    import re as _re

    def docx_number_suffix(number):
        return "." if _re.match(r"^[\dA-Z.]+$", number) else " —"

    def add_text_with_terms(paragraph, text, defs):
        """DOCX cross-reference simplification: defined terms render in
        bold (visually distinct, signalling "see Definitions") rather
        than as clickable internal hyperlinks. Word's internal-bookmark
        hyperlink mechanism needs the same raw-OOXML approach used for
        the letterhead background, and the marginal value over a real,
        working PDF (which does have internal links) plus a real,
        working website (real anchor links) didn't justify that added
        fragility here — an honest, documented scope decision, not an
        oversight."""
        if not defs:
            paragraph.add_run(text)
            return
        ordered = sorted(defs, key=lambda dd: -len(dd["term"]))
        pattern = _re.compile(r"\b(" + "|".join(_re.escape(dd["term"]) for dd in ordered) + r")\b")
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            run = paragraph.add_run(m.group(0))
            run.bold = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    for sec in sections:
        level = 1 if sec["level"] == 1 else 2
        d.add_heading(f"{sec['number']}{docx_number_suffix(sec['number'])} {sec['heading']}", level=level)
        linkable_defs = [] if sec["id"] == "definitions" else definitions
        for node in sec["content"]:
            if node["type"] == "paragraph":
                p = d.add_paragraph()
                add_text_with_terms(p, node["text"], linkable_defs)
            elif node["type"] == "subheading":
                p = d.add_paragraph()
                r = p.add_run(node["text"]); r.bold = True
            elif node["type"] == "list":
                for item in node["items"]:
                    p = d.add_paragraph(style="List Bullet")
                    add_text_with_terms(p, item, linkable_defs)
            elif node["type"] == "table":
                t = d.add_table(rows=1, cols=len(node["headers"]))
                for i, hcell in enumerate(node["headers"]):
                    t.rows[0].cells[i].text = hcell
                    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                for row_data in node["rows"]:
                    row = t.add_row().cells
                    for i, cell_val in enumerate(row_data):
                        row[i].text = cell_val
                d.add_paragraph()
        if sec["id"] == "definitions" and definitions:
            for defn in definitions:
                p = d.add_paragraph()
                r1 = p.add_run(f"“{defn['term']}” "); r1.bold = True
                p.add_run(defn["definition"])

    d.add_paragraph()
    footer_note = d.add_paragraph(
        f"© {fmt_date(control['effectiveDate'])[-4:]} Ordift Studios. All rights reserved. "
        f"Document {control['documentCode']}, Version {control['version']}."
    )
    footer_note.runs[0].font.size = Pt(8)

    d.save(out_path)


# ---------------------------------------------------------------------
# HTML (standalone, self-contained)
# ---------------------------------------------------------------------
def generate_html(doc, out_path):
    import re as _re

    control = doc["control"]
    sections = doc["sections"]
    definitions = doc["definitions"]

    def esc(s):
        return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def number_suffix(number):
        return "." if _re.match(r"^[\dA-Z.]+$", number) else " —"

    def link_terms(text, defs):
        if not defs:
            return esc(text)
        ordered = sorted(defs, key=lambda dd: -len(dd["term"]))
        pattern = _re.compile(r"\b(" + "|".join(_re.escape(dd["term"]) for dd in ordered) + r")\b")
        by_term = {dd["term"]: dd["id"] for dd in ordered}
        # Escape first, then link — safe here because none of the
        # defined terms contain characters esc() rewrites.
        escaped = esc(text)
        return pattern.sub(lambda m: f'<a href="#definitions-{by_term[m.group(0)]}" class="term-link">{m.group(0)}</a>', escaped)

    def render_node(node, linkable_defs):
        if node["type"] == "paragraph":
            return f"<p>{link_terms(node['text'], linkable_defs)}</p>"
        if node["type"] == "subheading":
            return f"<h4>{esc(node['text'])}</h4>"
        if node["type"] == "list":
            tag = "ol" if node.get("ordered") else "ul"
            items = "".join(f"<li>{link_terms(i, linkable_defs)}</li>" for i in node["items"])
            return f"<{tag}>{items}</{tag}>"
        if node["type"] == "table":
            headers = "".join(f"<th>{esc(h)}</th>" for h in node["headers"])
            rows = "".join("<tr>" + "".join(f"<td>{esc(c)}</td>" for c in r) + "</tr>" for r in node["rows"])
            return f"<table><thead><tr>{headers}</tr></thead><tbody>{rows}</tbody></table>"
        if node["type"] == "divider":
            return "<hr>"
        return ""

    toc_items = "".join(
        f'<li><a href="#{s["id"]}">{esc(s["number"])}{number_suffix(s["number"])} {esc(s["heading"])}</a></li>'
        for s in sections if s["level"] == 1
    )

    body_sections = []
    for s in sections:
        tag = "h2" if s["level"] == 1 else "h3"
        linkable_defs = [] if s["id"] == "definitions" else definitions
        body_sections.append(f'<section id="{s["id"]}"><{tag}>{esc(s["number"])}{number_suffix(s["number"])} {esc(s["heading"])}</{tag}>')
        body_sections.append("".join(render_node(n, linkable_defs) for n in s["content"]))
        if s["id"] == "definitions" and definitions:
            dl = "".join(f'<dt id="definitions-{d["id"]}">“{esc(d["term"])}”</dt><dd>{esc(d["definition"])}</dd>' for d in definitions)
            body_sections.append(f"<dl>{dl}</dl>")
        body_sections.append("</section>")

    related = "".join(
        f'<li>{esc((rd["code"] + " ") if rd["code"] else "")}{esc(rd["title"])}</li>' for rd in control["relatedDocuments"]
    )

    change_log_rows = "".join(
        f"<tr><td>{esc(e['version'])}</td><td>{fmt_date(e['date'])}</td><td>{esc(e['description'])}</td><td>{esc(e['author'])}</td></tr>"
        for e in control["changeLog"]
    )

    html = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(control['documentTitle'])} — Ordift Studios</title>
<meta name="description" content="{esc(control['documentTitle'])} ({esc(control['documentCode'])}, Version {esc(control['version'])}) — {esc(control['publicationSeries'])}.">
<style>
  :root {{ color-scheme: light; }}
  body {{ font-family: -apple-system, "Segoe UI", Helvetica, Arial, sans-serif; color: #1A1D24; max-width: 780px; margin: 0 auto; padding: 3rem 1.5rem 5rem; line-height: 1.6; }}
  h1 {{ font-size: 2rem; margin-bottom: 0.25rem; }}
  h2 {{ font-size: 1.375rem; margin-top: 2.5rem; border-top: 1px solid #E5E3DE; padding-top: 1.5rem; }}
  h3 {{ font-size: 1.125rem; margin-top: 1.5rem; }}
  h4 {{ font-size: 1rem; margin-top: 1rem; }}
  .eyebrow {{ text-transform: uppercase; letter-spacing: 0.15em; font-size: 0.75rem; font-weight: 700; color: #8E7736; }}
  .meta {{ color: #5B5F6B; font-size: 0.875rem; margin-bottom: 2rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 0.9rem; }}
  th, td {{ border: 1px solid #E5E3DE; padding: 0.5rem 0.75rem; text-align: left; vertical-align: top; }}
  th {{ background: #F7F5F1; }}
  dt {{ font-weight: 700; margin-top: 0.75rem; scroll-margin-top: 1.5rem; }}
  dd {{ margin-left: 0; color: #5B5F6B; }}
  a.term-link {{ color: inherit; text-decoration: underline dotted; text-underline-offset: 2px; }}
  a.term-link:hover {{ color: #8E7736; }}
  nav.toc {{ background: #F7F5F1; border-radius: 0.75rem; padding: 1rem 1.5rem; margin: 2rem 0; }}
  nav.toc ol {{ margin: 0; padding-left: 1.25rem; }}
  .doc-control {{ border: 1px solid #E5E3DE; border-radius: 0.75rem; padding: 1rem 1.5rem; margin: 2rem 0; }}
  .doc-control dt {{ color: #5B5F6B; font-size: 0.8rem; margin-top: 0.6rem; }}
  .doc-control dd {{ margin-left: 0; }}
  footer.copyright {{ margin-top: 3rem; border-top: 1px solid #E5E3DE; padding-top: 1rem; font-size: 0.8rem; color: #5B5F6B; }}
</style>
</head>
<body>
<p class="eyebrow">{esc(control['publicationSeries'])}</p>
<h1>{esc(control['documentTitle'])}</h1>
<p class="meta">{esc(control['documentCode'])} · Version {esc(control['version'])} · Effective {fmt_date(control['effectiveDate'])}</p>

<section class="doc-control" aria-labelledby="doc-control-heading">
  <h2 id="doc-control-heading" style="margin-top:0;border-top:none;padding-top:0;">Document Control</h2>
  <dl>
    <dt>Status</dt><dd>{control['status'].capitalize()}</dd>
    <dt>Classification</dt><dd>{control['classification'].capitalize()}</dd>
    <dt>Last Updated</dt><dd>{fmt_date(control['lastUpdated'])}</dd>
    <dt>Review Cycle</dt><dd>{esc(control['reviewCycle'])}</dd>
    <dt>Document Owner</dt><dd>{esc(control['documentOwner'])}</dd>
    <dt>Prepared By</dt><dd>{esc(control['preparedBy'])}</dd>
    <dt>Approved By</dt><dd>{esc(control['approvedBy'])}</dd>
    <dt>Related Documents</dt><dd><ul>{related}</ul></dd>
  </dl>
  <p><strong>Change Log</strong></p>
  <table>
    <thead><tr><th>Version</th><th>Date</th><th>Description</th><th>Author</th></tr></thead>
    <tbody>{change_log_rows}</tbody>
  </table>
  <p>{esc(control['controlledDocumentNotice'])}</p>
</section>

<nav class="toc" aria-label="Table of contents">
  <strong>Table of Contents</strong>
  <ol>{toc_items}</ol>
</nav>

<article>
{''.join(body_sections)}
</article>

<footer class="copyright">
  &copy; {fmt_date(control['effectiveDate'])[-4:]} Ordift Studios. All rights reserved. Document {esc(control['documentCode'])}, Version {esc(control['version'])}.
</footer>
</body>
</html>
"""
    with open(out_path, "w") as f:
        f.write(html)


# ---------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------
def generate_markdown(doc, out_path):
    import re as _re

    control = doc["control"]
    sections = doc["sections"]
    definitions = doc["definitions"]

    def number_suffix(number):
        return "." if _re.match(r"^[\dA-Z.]+$", number) else " —"

    def link_terms(text, defs):
        if not defs:
            return text
        ordered = sorted(defs, key=lambda dd: -len(dd["term"]))
        pattern = _re.compile(r"\b(" + "|".join(_re.escape(dd["term"]) for dd in ordered) + r")\b")
        by_term = {dd["term"]: dd["id"] for dd in ordered}
        return pattern.sub(lambda m: f"[{m.group(0)}](#definitions-{by_term[m.group(0)]})", text)

    lines = []
    lines.append(f"# {control['documentTitle']}")
    lines.append("")
    lines.append(f"*{control['publicationSeries']}*")
    lines.append("")
    lines.append(f"**{control['documentCode']}** · Version {control['version']} · Effective {fmt_date(control['effectiveDate'])}")
    lines.append("")
    lines.append("## Document Control")
    lines.append("")
    lines.append("| Field | Value |")
    lines.append("|---|---|")
    dc_rows = [
        ("Document Title", control["documentTitle"]),
        ("Document Code", control["documentCode"]),
        ("Publication Series", control["publicationSeries"]),
        ("Version", control["version"]),
        ("Status", control["status"].capitalize()),
        ("Classification", control["classification"].capitalize()),
        ("Effective Date", fmt_date(control["effectiveDate"])),
        ("Last Updated", fmt_date(control["lastUpdated"])),
        ("Review Cycle", control["reviewCycle"]),
        ("Document Owner", control["documentOwner"]),
        ("Prepared By", control["preparedBy"]),
        ("Approved By", control["approvedBy"]),
    ]
    for k, v in dc_rows:
        lines.append(f"| {k} | {v} |")
    related = "; ".join(f"{(rd['code'] + ' ') if rd['code'] else ''}{rd['title']}" for rd in control["relatedDocuments"])
    lines.append(f"| Related Documents | {related} |")
    lines.append("")
    lines.append("**Change Log**")
    lines.append("")
    lines.append("| Version | Date | Description | Author |")
    lines.append("|---|---|---|---|")
    for entry in control["changeLog"]:
        lines.append(f"| {entry['version']} | {fmt_date(entry['date'])} | {entry['description']} | {entry['author']} |")
    lines.append("")
    lines.append(f"> {control['controlledDocumentNotice']}")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("")
    for s in sections:
        if s["level"] == 1:
            lines.append(f"{s['number']}{number_suffix(s['number'])} [{s['heading']}](#{s['id']})")
    lines.append("")

    for s in sections:
        heading_prefix = "##" if s["level"] == 1 else "###"
        linkable_defs = [] if s["id"] == "definitions" else definitions
        lines.append(f'<a id="{s["id"]}"></a>')
        lines.append(f"{heading_prefix} {s['number']}{number_suffix(s['number'])} {s['heading']}")
        lines.append("")
        for node in s["content"]:
            if node["type"] == "paragraph":
                lines.append(link_terms(node["text"], linkable_defs))
                lines.append("")
            elif node["type"] == "subheading":
                lines.append(f"**{node['text']}**")
                lines.append("")
            elif node["type"] == "list":
                for i, item in enumerate(node["items"]):
                    marker = f"{i + 1}." if node.get("ordered") else "-"
                    lines.append(f"{marker} {link_terms(item, linkable_defs)}")
                lines.append("")
            elif node["type"] == "table":
                lines.append("| " + " | ".join(node["headers"]) + " |")
                lines.append("|" + "|".join(["---"] * len(node["headers"])) + "|")
                for row in node["rows"]:
                    lines.append("| " + " | ".join(c.replace("\n", " ") for c in row) + " |")
                lines.append("")
        if s["id"] == "definitions" and definitions:
            for d in definitions:
                lines.append(f'<a id="definitions-{d["id"]}"></a>**“{d["term"]}”** {d["definition"]}')
                lines.append("")

    lines.append("---")
    lines.append("")
    lines.append(f"© {fmt_date(control['effectiveDate'])[-4:]} Ordift Studios. All rights reserved. Document {control['documentCode']}, Version {control['version']}.")

    with open(out_path, "w") as f:
        f.write("\n".join(lines))


def main():
    if len(sys.argv) != 3:
        print("Usage: python3 scripts/generateLegalPublication.py <json-path> <out-dir>")
        sys.exit(1)
    json_path, out_dir = sys.argv[1], sys.argv[2]
    os.makedirs(out_dir, exist_ok=True)
    doc = load_doc(json_path)
    code_lower = doc["control"]["documentCode"].lower()

    generate_pdf(doc, os.path.join(out_dir, f"{code_lower}.pdf"))
    print(f"Generated {code_lower}.pdf")
    generate_docx(doc, os.path.join(out_dir, f"{code_lower}.docx"))
    print(f"Generated {code_lower}.docx")
    generate_html(doc, os.path.join(out_dir, f"{code_lower}.html"))
    print(f"Generated {code_lower}.html")
    generate_markdown(doc, os.path.join(out_dir, f"{code_lower}.md"))
    print(f"Generated {code_lower}.md")


if __name__ == "__main__":
    main()
